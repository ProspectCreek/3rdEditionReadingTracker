import sys
import os
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QSplitter, QListWidget, QListWidgetItem,
    QLabel, QTextEdit, QFrame, QPushButton, QMessageBox, QFileDialog
)
from PySide6.QtCore import Qt, Signal, QTimer
from PySide6.QtGui import QColor, QIcon, QPixmap, QPainter
from PySide6.QtSvg import QSvgRenderer

try:
    from docx import Document
except ImportError:
    Document = None


class AnnotatedBibTab(QWidget):
    """
    Tab for managing Annotated Bibliography entries.
    Layout: Left (Source List), Right (Fields: Citation, Description, Analysis, Applicability)
    """

    def __init__(self, db, project_id, parent=None):
        super().__init__(parent)
        self.db = db
        self.project_id = project_id
        self.current_reading_id = None
        self._ignore_changes = False

        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # --- Header / Export Bar ---
        header_layout = QHBoxLayout()
        header_layout.setContentsMargins(8, 8, 8, 0)

        title_label = QLabel("Annotated Bibliography")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; color: #333;")

        self.export_btn = QPushButton("Export Annotated Bibliography")
        self.export_btn.clicked.connect(self.export_data)

        header_layout.addWidget(title_label)
        header_layout.addStretch()
        header_layout.addWidget(self.export_btn)

        main_layout.addLayout(header_layout)

        # --- Main Splitter ---
        self.splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(self.splitter)

        # --- Left Panel: Sources ---
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(4, 4, 4, 4)

        left_layout.addWidget(QLabel("<b>Readings</b>"))
        self.source_list = QListWidget()
        self.source_list.currentItemChanged.connect(self.on_source_selected)
        left_layout.addWidget(self.source_list)

        # --- Right Panel: Editor ---
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(4, 4, 4, 4)

        # Scroll area for fields could be good, but fixed fields fit well if not huge
        # Let's use a simple layout of vertically stacked labels and text edits

        # 1. Citation
        self.citation_edit = self._create_field(right_layout, "Citation:", 80)

        # 2. Description
        self.desc_edit = self._create_field(right_layout, "Description – What the source says:", 100)

        # 3. Analysis
        self.analysis_edit = self._create_field(right_layout, "Analysis – Evaluation / strengths / limitations:", 100)

        # 4. Applicability
        self.applicability_edit = self._create_field(right_layout,
                                                     "Applicability – How this source relates to my research:", 100)

        # Connect changes
        for edit in [self.citation_edit, self.desc_edit, self.analysis_edit, self.applicability_edit]:
            edit.textChanged.connect(self.save_current_data)

        self.splitter.addWidget(left_panel)
        self.splitter.addWidget(right_panel)
        self.splitter.setSizes([300, 700])

        self.load_sources()

    def _create_field(self, layout, label_text, height):
        lbl = QLabel(label_text)
        lbl.setStyleSheet("font-weight: bold; color: #555; margin-top: 8px;")
        layout.addWidget(lbl)

        edit = QTextEdit()
        edit.setMinimumHeight(height)
        # edit.setMaximumHeight(height + 50) # Optional max height
        layout.addWidget(edit)
        return edit

    def load_sources(self):
        """Loads readings and their annotation status."""
        self.source_list.clear()
        entries = self.db.get_annotated_bib_entries(self.project_id)

        for entry in entries:
            nickname = entry.get('nickname') or entry.get('title', 'Untitled')

            # Determine Status
            status = entry.get('status', 'Not Started')

            # Simple icon or text color for status
            if status == "Completed":
                icon_text = "🟢"
            elif status == "In Process":
                icon_text = "🟡"
            else:
                icon_text = "⚪"  # Not Started

            display_text = f"{icon_text}  {nickname}"

            item = QListWidgetItem(display_text)
            item.setData(Qt.ItemDataRole.UserRole, entry['reading_id'])
            item.setData(Qt.ItemDataRole.UserRole + 1, entry)  # Store full data

            self.source_list.addItem(item)

    def on_source_selected(self, current, previous):
        # Save previous if needed (though textChanged handles it)

        if not current:
            self.clear_fields()
            self.current_reading_id = None
            return

        self._ignore_changes = True

        reading_id = current.data(Qt.ItemDataRole.UserRole)
        data = current.data(Qt.ItemDataRole.UserRole + 1)
        self.current_reading_id = reading_id

        self.citation_edit.setPlainText(data.get('citation_text') or "")
        self.desc_edit.setPlainText(data.get('description') or "")
        self.analysis_edit.setPlainText(data.get('analysis') or "")
        self.applicability_edit.setPlainText(data.get('applicability') or "")

        self._ignore_changes = False

    def clear_fields(self):
        self._ignore_changes = True
        self.citation_edit.clear()
        self.desc_edit.clear()
        self.analysis_edit.clear()
        self.applicability_edit.clear()
        self._ignore_changes = False

    def save_current_data(self):
        if self._ignore_changes or not self.current_reading_id:
            return

        # Gather data
        citation = self.citation_edit.toPlainText().strip()
        desc = self.desc_edit.toPlainText().strip()
        analysis = self.analysis_edit.toPlainText().strip()
        app = self.applicability_edit.toPlainText().strip()

        # Determine status
        if citation and desc and analysis and app:
            status = "Completed"
        elif citation or desc or analysis or app:
            status = "In Process"
        else:
            status = "Not Started"

        data = {
            "citation_text": citation,
            "description": desc,
            "analysis": analysis,
            "applicability": app,
            "status": status
        }

        self.db.update_annotated_bib_entry(self.current_reading_id, data)

        # Update list item status
        current_item = self.source_list.currentItem()
        if current_item:
            # Update stored data
            old_data = current_item.data(Qt.ItemDataRole.UserRole + 1)
            old_data.update(data)
            current_item.setData(Qt.ItemDataRole.UserRole + 1, old_data)

            # Update display text
            nickname = old_data.get('nickname') or old_data.get('title', 'Untitled')
            if status == "Completed":
                icon_text = "🟢"
            elif status == "In Process":
                icon_text = "🟡"
            else:
                icon_text = "⚪"
            current_item.setText(f"{icon_text}  {nickname}")

    def export_data(self):
        filename, _ = QFileDialog.getSaveFileName(self, "Export Annotated Bibliography", "Annotated_Bib.docx",
                                                  "Word Document (*.docx);;Text File (*.txt)")
        if not filename:
            return

        try:
            if filename.endswith(".docx") and Document:
                self._export_docx(filename)
            else:
                self._export_txt(filename)

            QMessageBox.information(self, "Success", f"Exported to {filename}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Export failed: {e}")

    def _export_docx(self, filename):
        doc = Document()
        doc.add_heading("Annotated Bibliography", 0)

        entries = self.db.get_annotated_bib_entries(self.project_id)

        for entry in entries:
            cit = (entry.get('citation_text') or "").strip()
            desc = (entry.get('description') or "").strip()
            ana = (entry.get('analysis') or "").strip()
            app = (entry.get('applicability') or "").strip()

            # Skip empty entries?
            if not (cit or desc or ana or app):
                continue

            if cit:
                p = doc.add_paragraph(cit)
                # Hanging indent is typical for bibs, but standard style is fine for now.

            if desc:
                doc.add_paragraph(desc)
            if ana:
                doc.add_paragraph(ana)
            if app:
                doc.add_paragraph(app)

            doc.add_paragraph("")  # Spacer

        doc.save(filename)

    def _export_txt(self, filename):
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("Annotated Bibliography\n")
            f.write("======================\n\n")

            entries = self.db.get_annotated_bib_entries(self.project_id)
            for entry in entries:
                cit = (entry.get('citation_text') or "").strip()
                desc = (entry.get('description') or "").strip()
                ana = (entry.get('analysis') or "").strip()
                app = (entry.get('applicability') or "").strip()

                if not (cit or desc or ana or app):
                    continue

                if cit: f.write(f"{cit}\n")
                if desc: f.write(f"{desc}\n")
                if ana: f.write(f"{ana}\n")
                if app: f.write(f"{app}\n")
                f.write("\n")