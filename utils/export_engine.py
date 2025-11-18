# utils/export_engine.py
import os
import re
import traceback
from bs4 import BeautifulSoup

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ExportEngine Warning: 'python-docx' not installed. DOCX export will fail.")
    Document = None


def html_to_text(html):
    """Strips HTML tags to get plain text."""
    if not html:
        return ""
    try:
        soup = BeautifulSoup(html, 'html.parser')
        # Replace <p> and <br> with newlines for basic formatting
        for tag in soup.find_all(['p', 'br', 'li', 'h1', 'h2', 'h3']):
            tag.append('\n')
        return soup.get_text().strip()
    except Exception:
        return ""  # Return empty on any parsing error


def clean_html_for_docx(html):
    """A simple cleaner for basic HTML tags before docx conversion."""
    if not html:
        return []

    # Use BeautifulSoup to parse and iterate
    soup = BeautifulSoup(html, 'html.parser')

    paragraphs = []

    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'li']):
        text = element.get_text().strip()
        style = 'BodyText'

        if element.name in ['h1', 'h2', 'h3']:
            style = f"Heading {element.name[1]}"
        elif element.name == 'li':
            # Crude way to show list indentation
            text = f"• {text}"
            style = 'ListParagraph'

        if text:  # Only add if there is text
            paragraphs.append({'text': text, 'style': style})

    if not paragraphs and soup.get_text().strip():
        # Fallback for plain text not in a known tag
        paragraphs.append({'text': soup.get_text().strip(), 'style': 'BodyText'})

    return paragraphs


class ExportEngine:
    """
    Handles fetching data and compiling it into a single file
    in the specified format (HTML, DOCX, or TXT).
    """

    def __init__(self, db, project_id):
        self.db = db
        self.project_id = project_id
        self.project_details = self.db.get_item_details(self.project_id)
        self.project_name = self.project_details.get('name', 'Untitled Project')

        # Set default font
        self.font_name = "Times New Roman"

        # Caches
        self.readings_cache = None
        self.terminology_cache = None
        self.propositions_cache = None
        self.todo_cache = None
        self.rubric_cache = None

    def export_to_file(self, file_path, config):
        """Public method to generate and save the export file."""
        file_format = config.get("format", "html")
        components = config.get("components", [])

        if file_format == "html":
            content = self.generate_html(components)
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
        elif file_format == "txt":
            content = self.generate_txt(components)
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
        elif file_format == "docx":
            if Document is None:
                raise ImportError("The 'python-docx' library is required for .docx export. Please install it.")
            doc = self.generate_docx(components)
            doc.save(file_path)
        else:
            raise ValueError(f"Unsupported format: {file_format}")

    # --- Data Fetchers (with Caching) ---

    def get_readings(self):
        if self.readings_cache is None:
            self.readings_cache = self.db.get_readings(self.project_id)
        return self.readings_cache

    def get_reading_details(self, reading_id):
        """Gets just the main details for a reading."""
        return self.db.get_reading_details(reading_id)

    def get_reading_outline_data(self, reading_id, parent_id=None, indent=0):
        """Recursively fetches the full outline with notes."""
        items = self.db.get_reading_outline(reading_id, parent_id)
        export_items = []
        for item in items:
            item_data = dict(item)
            item_data['notes_html'] = self.db.get_outline_section_notes(item['id'])
            item_data['indent'] = indent
            export_items.append(item_data)
            export_items.extend(self.get_reading_outline_data(reading_id, item['id'], indent + 1))
        return export_items

    def get_reading_dqs(self, reading_id):
        return self.db.get_driving_questions(reading_id, parent_id=True)  # Get all

    def get_reading_key_terms(self, reading_id):
        return self.db.get_reading_key_terms(reading_id)

    def get_reading_leading_propositions(self, reading_id):
        return self.db.get_reading_propositions_simple(reading_id)

    def get_reading_arguments(self, reading_id):
        return self.db.get_reading_arguments(reading_id)

    def get_reading_theories(self, reading_id):
        return self.db.get_reading_theories(reading_id)

    def get_terminology(self):
        if self.terminology_cache is None:
            self.terminology_cache = self.db.get_project_terminology(self.project_id)
        return self.terminology_cache

    def get_propositions(self):
        if self.propositions_cache is None:
            self.propositions_cache = self.db.get_project_propositions(self.project_id)
        return self.propositions_cache

    def get_todo_list(self):
        if self.todo_cache is None:
            self.todo_cache = self.db.get_project_todo_items(self.project_id)
        return self.todo_cache

    def get_rubric(self):
        if self.rubric_cache is None:
            self.rubric_cache = self.db.get_rubric_components(self.project_id)
        return self.rubric_cache

    # --- Main Generator Functions ---

    def generate_html(self, components):
        """Generates a single HTML string."""
        body = f"<h1>{self.project_name}</h1>"

        for comp in components:
            key = comp['key']
            title = comp['title']

            # --- NEW: Parse key ---
            reading_id = None
            if "_" in key:
                parts = key.split('_')
                try:
                    reading_id = int(parts[-1])
                    key = "_".join(parts[:-1])  # Reconstruct key, e.g., "reading_driving_questions"
                except ValueError:
                    pass  # Not a reading component
            # --- END NEW ---

            if key.startswith("reading_"):
                # Reading sub-components get smaller headers
                body += f"<h3>{title}</h3>"
            else:
                # Project-level components
                body += f"<h2>{title}</h2>"

            # --- Project-level components ---
            if key in self.project_details:  # Catches all simple text fields
                body += self.project_details.get(key, "") or "<p><i>(No content)</i></p>"
            elif key == "assignment_rubric":
                body += self._get_rubric_html()
            elif key == "synthesis_terminology":
                body += self._get_terminology_html()
            elif key == "synthesis_propositions":
                body += self._get_propositions_html()
            elif key == "todo_list":
                body += self._get_todo_html()

            # --- Reading-level components ---
            elif key == "reading_header":
                body += self._get_reading_header_html(reading_id)
            elif key == "reading_outline":
                body += self._get_reading_outline_html(reading_id)
            elif key == "reading_driving_questions":
                body += self._get_reading_dqs_html(readingid)
            elif key == "reading_key_terms":
                body += self._get_reading_key_terms_html(reading_id)
            elif key == "reading_leading_propositions":
                body += self._get_reading_leading_propositions_html(reading_id)
            elif key == "reading_arguments":
                body += self._get_reading_arguments_html(reading_id)
            elif key == "reading_theories":
                body += self._get_reading_theories_html(reading_id)
            elif key == "reading_unity":
                body += self._get_reading_unity_html(reading_id)

            else:
                body += f"<p><i>(Component '{key}' not implemented)</i></p>"

            if not key.startswith("reading_"):
                body += "<hr>"  # Add spacer after top-level components

        return f"""
        <html>
            <head>
                <meta charset="utf-8">
                <title>{self.project_name}</title>
                <style>
                    body {{ font-family: "{self.font_name}", serif; line-height: 1.5; padding: 2em; }}
                    h1, h2, h3, h4 {{ font-family: "{self.font_name}", serif; }}
                    h1 {{ font-size: 2.5em; }}
                    h2 {{ font-size: 2em; border-bottom: 2px solid #333; padding-bottom: 5px; margin-top: 1.5em; }}
                    h3 {{ font-size: 1.5em; margin-top: 1.5em; }}
                    h4 {{ font-size: 1.2em; margin-top: 1em; }}
                    hr {{ border: 0; border-top: 1px solid #ccc; margin: 2em 0; }}
                    .component-box {{
                        background: #fdfdfd;
                        border: 1px solid #eee;
                        padding: 1.5em;
                        border-radius: 5px;
                        margin-bottom: 1.5em;
                    }}
                    .component-box h4 {{ margin-top: 0.5em; }}
                    blockquote {{ border-left: 3px solid #ccc; padding-left: 1em; margin-left: 1em; }}
                    .outline-item {{ margin-left: 2em; }}
                    .outline-notes {{ background: #f9f9f9; border: 1px solid #eee; padding: 1em; margin-top: 0.5em; }}
                    .todo-item {{ list-style-type: none; margin-left: 1em; }}
                </style>
            </head>
            <body>{body}</body>
        </html>
        """

    def generate_txt(self, components):
        """Generates a single plain text string."""
        output = f"{self.project_name}\n"
        output += "=" * len(self.project_name) + "\n\n"

        for comp in components:
            key = comp['key']
            title = comp['title']

            # --- NEW: Parse key ---
            reading_id = None
            if "_" in key:
                parts = key.split('_')
                try:
                    reading_id = int(parts[-1])
                    key = "_".join(parts[:-1])
                except ValueError:
                    pass
            # --- END NEW ---

            # --- Project-level components ---
            if not key.startswith("reading_"):
                output += f"\n\n=======================================\n"
                output += f"{title.upper()}\n"
                output += "=======================================\n\n"

            if key in self.project_details:  # Simple text fields
                output += html_to_text(self.project_details.get(key, "")) or "(No content)"
            elif key == "assignment_rubric":
                output += self._get_rubric_txt()
            elif key == "synthesis_terminology":
                output += self._get_terminology_txt()
            elif key == "synthesis_propositions":
                output += self._get_propositions_txt()
            elif key == "todo_list":
                output += self._get_todo_txt()

            # --- Reading-level components ---
            elif key == "reading_header":
                output += self._get_reading_header_txt(reading_id)
            elif key == "reading_outline":
                output += f"\n--- {title} ---\n\n"
                output += self._get_reading_outline_txt(reading_id)
            elif key == "reading_driving_questions":
                output += f"\n--- {title} ---\n\n"
                output += self._get_reading_dqs_txt(reading_id)
            elif key == "reading_key_terms":
                output += f"\n--- {title} ---\n\n"
                output += self._get_reading_key_terms_txt(reading_id)
            elif key == "reading_leading_propositions":
                output += f"\n--- {title} ---\n\n"
                output += self._get_reading_leading_propositions_txt(reading_id)
            elif key == "reading_arguments":
                output += f"\n--- {title} ---\n\n"
                output += self._get_reading_arguments_txt(reading_id)
            elif key == "reading_theories":
                output += f"\n--- {title} ---\n\n"
                output += self._get_reading_theories_txt(reading_id)
            elif key == "reading_unity":
                output += f"\n--- {title} ---\n\n"
                output += self._get_reading_unity_txt(reading_id)
            else:
                output += f"(Component '{key}' not implemented)"

            output += "\n\n"

        return output

    def generate_docx(self, components):
        """Generates a .docx Document object."""
        if Document is None:
            return None

        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(12)

        # Set heading fonts
        for i in range(1, 5):
            # --- FIX: Access styles by name using indexing ---
            try:
                style = doc.styles[f"Heading {i}"]
                style.font.name = self.font_name
                style.font.bold = True
            except KeyError:
                print(f"Warning: 'Heading {i}' style not found in docx template.")
            # --- END FIX ---

        doc.add_heading(self.project_name, level=0)

        for comp in components:
            key = comp['key']
            title = comp['title']

            # --- NEW: Parse key ---
            reading_id = None
            if "_" in key:
                parts = key.split('_')
                try:
                    reading_id = int(parts[-1])
                    key = "_".join(parts[:-1])
                except ValueError:
                    pass
            # --- END NEW ---

            if not key.startswith("reading_"):
                doc.add_heading(title, level=1)

            if key in self.project_details:  # Simple text fields
                html_content = self.project_details.get(key, "")
                self._add_html_to_docx(doc, html_content)
            elif key == "assignment_rubric":
                self._add_rubric_to_docx(doc)
            elif key == "synthesis_terminology":
                self._add_terminology_to_docx(doc)
            elif key == "synthesis_propositions":
                self._add_propositions_to_docx(doc)
            elif key == "todo_list":
                self._add_todo_to_docx(doc)

            # --- Reading-level components ---
            elif key == "reading_header":
                self._add_reading_header_to_docx(doc, reading_id, title)
            elif key == "reading_outline":
                doc.add_heading(title, level=2)
                self._add_reading_outline_to_docx(doc, reading_id)
            elif key == "reading_driving_questions":
                doc.add_heading(title, level=2)
                self._add_reading_dqs_to_docx(doc, reading_id)
            elif key == "reading_key_terms":
                doc.add_heading(title, level=2)
                self._add_reading_key_terms_to_docx(doc, reading_id)
            elif key == "reading_leading_propositions":
                doc.add_heading(title, level=2)
                self._add_reading_leading_propositions_to_docx(doc, reading_id)
            elif key == "reading_arguments":
                doc.add_heading(title, level=2)
                self._add_reading_arguments_to_docx(doc, reading_id)
            elif key == "reading_theories":
                doc.add_heading(title, level=2)
                self._add_reading_theories_to_docx(doc, reading_id)
            elif key == "reading_unity":
                doc.add_heading(title, level=2)
                self._add_reading_unity_to_docx(doc, reading_id)

            else:
                doc.add_paragraph(f"(Component '{key}' not implemented)")

            if not key.startswith("reading_"):
                doc.add_page_break()

        return doc

    # --- DOCX Helper Methods ---

    def _add_html_to_docx(self, doc, html):
        """Adds basic HTML content to a DOCX document."""
        if not html:
            doc.add_paragraph("(No content)")
            return

        paragraphs = clean_html_for_docx(html)
        if not paragraphs:
            doc.add_paragraph("(No content)")
            return

        for p in paragraphs:
            # --- FIX: Use built-in styles if they exist ---
            try:
                doc.add_paragraph(p['text'], style=p['style'])
            except KeyError:
                # Fallback if style doesn't exist
                doc.add_paragraph(p['text'], style='BodyText')
            # --- END FIX ---

    def _add_rubric_to_docx(self, doc):
        rubric = self.get_rubric()
        if not rubric:
            doc.add_paragraph("(No rubric components)")
            return
        for item in rubric:
            prefix = "[x] " if item['is_checked'] else "[ ] "
            doc.add_paragraph(f"{prefix}{item['component_text']}", style='ListParagraph')

    def _add_todo_to_docx(self, doc):
        items = self.get_todo_list()
        if not items:
            doc.add_paragraph("(No to-do items)")
            return
        for item in items:
            prefix = "[x] " if item['is_checked'] else "[ ] "
            doc.add_paragraph(f"{prefix}{item['display_name']}", style='ListParagraph')

    def _add_reading_header_to_docx(self, doc, reading_id, title):
        doc.add_heading(title, level=1)
        try:
            reading = self.get_reading_details(reading_id)
            if reading.get('author'):
                doc.add_paragraph(f"Author: {reading['author']}")
        except Exception as e:
            doc.add_paragraph(f"(Error loading reading header: {e})")

    def _add_reading_outline_to_docx(self, doc, reading_id):
        try:
            outline = self.get_reading_outline_data(reading_id)
            for item in outline:
                indent = item.get('indent', 0)
                p = doc.add_paragraph(f"{'  ' * indent}• {item['section_title']}")
                p.paragraph_format.left_indent = Inches(0.25 * indent)

                notes_html = item.get('notes_html', '')
                if notes_html:
                    self._add_html_to_docx(doc, notes_html)

        except Exception as e:
            doc.add_paragraph(f"(Error processing outline: {e})")

    def _add_reading_dqs_to_docx(self, doc, reading_id):
        items = self.get_reading_dqs(reading_id)
        if not items:
            doc.add_paragraph("(No driving questions)")
            return
        for item in items:
            doc.add_paragraph(item.get('nickname') or item.get('question_text', '...'), style='Heading 4')
            self._add_html_to_docx(doc, item.get('why_question', ''))

    def _add_reading_key_terms_to_docx(self, doc, reading_id):
        items = self.get_reading_key_terms(reading_id)
        if not items:
            doc.add_paragraph("(No key terms)")
            return
        for item_summary in items:
            item = self.db.get_reading_key_term_details(item_summary['id'])
            doc.add_paragraph(item.get('term', '...'), style='Heading 4')
            self._add_html_to_docx(doc, item.get('definition', ''))

    def _add_reading_leading_propositions_to_docx(self, doc, reading_id):
        items = self.get_reading_leading_propositions(reading_id)
        if not items:
            doc.add_paragraph("(No leading propositions)")
            return
        for item_summary in items:
            item = self.db.get_reading_proposition_details(item_summary['id'])
            doc.add_paragraph(item.get('nickname') or item.get('proposition_text', '...'), style='Heading 4')
            self._add_html_to_docx(doc, item.get('proposition_text', ''))

    def _add_reading_arguments_to_docx(self, doc, reading_id):
        items = self.get_reading_arguments(reading_id)
        if not items:
            doc.add_paragraph("(No arguments)")
            return
        for item_summary in items:
            item = self.db.get_argument_details(item_summary['id'])
            doc.add_paragraph(item.get('claim_text', '...'), style='Heading 4')
            doc.add_paragraph(f"Because: {item.get('because_text', '...')}")

    def _add_reading_theories_to_docx(self, doc, reading_id):
        items = self.get_reading_theories(reading_id)
        if not items:
            doc.add_paragraph("(No theories)")
            return
        for item_summary in items:
            item = self.db.get_reading_theory_details(item_summary['id'])
            doc.add_paragraph(item.get('theory_name', '...'), style='Heading 4')
            self._add_html_to_docx(doc, item.get('description', ''))

    def _add_reading_unity_to_docx(self, doc, reading_id):
        reading = self.get_reading_details(reading_id)
        doc.add_paragraph(f"Kind of Work: {reading.get('unity_kind_of_work', 'N/A')}")
        self._add_html_to_docx(doc, reading.get('unity_html', ''))

    def _add_terminology_to_docx(self, doc):
        items = self.get_terminology()
        if not items:
            doc.add_paragraph("(No terminology added)")
            return
        for item in items:
            term = self.db.get_terminology_details(item['id'])
            doc.add_heading(term['term'], level=3)
            self._add_html_to_docx(doc, term.get('meaning', ''))

    def _add_propositions_to_docx(self, doc):
        items = self.get_propositions()
        if not items:
            doc.add_paragraph("(No propositions added)")
            return
        for item in items:
            prop = self.db.get_proposition_details(item['id'])
            doc.add_heading(prop['display_name'], level=3)
            self._add_html_to_docx(doc, prop.get('proposition_html', ''))

    # --- HTML Helper Methods ---

    def _get_rubric_html(self):
        rubric = self.get_rubric()
        if not rubric:
            return "<p><i>(No rubric components)</i></p>"
        html = "<ul>"
        for item in rubric:
            check = "☑" if item['is_checked'] else "☐"
            html += f"<li>{check} {item['component_text']}</li>"
        html += "</ul>"
        return html

    def _get_todo_html(self):
        items = self.get_todo_list()
        if not items:
            return "<p><i>(No to-do items)</i></p>"
        html = "<ul>"
        for item in items:
            check = "☑" if item['is_checked'] else "☐"
            html += f"<li class='todo-item'>{check} {item['display_name']}</li>"
        html += "</ul>"
        return html

    def _get_reading_header_html(self, reading_id):
        try:
            reading = self.get_reading_details(reading_id)
            html = ""
            if reading.get('author'):
                html += f"<p><b>Author:</b> {reading['author']}</p>"
            return html
        except Exception as e:
            return f"<p><i>(Error processing reading header: {e})</i></p>"

    def _get_reading_outline_html(self, reading_id):
        try:
            outline = self.get_reading_outline_data(reading_id)
            if not outline:
                html = "<p><i>(No outline for this reading)</i></p>"
                return html

            html = ""
            for item in outline:
                indent = item.get('indent', 0)
                html += f"<div class='outline-item' style='margin-left: {indent * 2}em;'>"
                html += f"<h4>{item['section_title']}</h4>"
                notes_html = item.get('notes_html', '')
                if notes_html:
                    html += f"<div class='outline-notes'>{notes_html}</div>"
                html += "</div>"
            return html
        except Exception as e:
            return f"<p><i>(Error processing outline: {e})</i></p>"

    def _get_reading_dqs_html(self, reading_id):
        items = self.get_reading_dqs(reading_id)
        if not items: return "<p><i>(No driving questions)</i></p>"
        html = ""
        for item in items:
            html += "<div class='component-box'>"
            html += f"<h4>{item.get('nickname') or item.get('question_text', '...')}</h4>"
            html += f"<div>{item.get('why_question', '')}</div>"
            html += "</div>"
        return html

    def _get_reading_key_terms_html(self, reading_id):
        items = self.get_reading_key_terms(reading_id)
        if not items: return "<p><i>(No key terms)</i></p>"
        html = ""
        for item_summary in items:
            item = self.db.get_reading_key_term_details(item_summary['id'])
            html += "<div class='component-box'>"
            html += f"<h4>{item.get('term', '...')}</h4>"
            html += f"<div>{item.get('definition', '')}</div>"
            html += "</div>"
        return html

    def _get_reading_leading_propositions_html(self, reading_id):
        items = self.get_reading_leading_propositions(reading_id)
        if not items: return "<p><i>(No leading propositions)</i></p>"
        html = ""
        for item_summary in items:
            item = self.db.get_reading_proposition_details(item_summary['id'])
            html += "<div class='component-box'>"
            html += f"<h4>{item.get('nickname') or item.get('proposition_text', '...')}</h4>"
            html += f"<div>{item.get('proposition_text', '')}</div>"
            html += "</div>"
        return html

    def _get_reading_arguments_html(self, reading_id):
        items = self.get_reading_arguments(reading_id)
        if not items: return "<p><i>(No arguments)</i></p>"
        html = ""
        for item_summary in items:
            item = self.db.get_argument_details(item_summary['id'])
            html += "<div class='component-box'>"
            html += f"<h4>{item.get('claim_text', '...')}</h4>"
            html += f"<p><b>Because:</b> {item.get('because_text', '...')}</p>"
            html += "</div>"
        return html

    def _get_reading_theories_html(self, reading_id):
        items = self.get_reading_theories(reading_id)
        if not items: return "<p><i>(No theories)</i></p>"
        html = ""
        for item_summary in items:
            item = self.db.get_reading_theory_details(item_summary['id'])
            html += "<div class='component-box'>"
            html += f"<h4>{item.get('theory_name', '...')}</h4>"
            html += f"<div>{item.get('description', '')}</div>"
            html += "</div>"
        return html

    def _get_reading_unity_html(self, reading_id):
        reading = self.get_reading_details(reading_id)
        html = "<div class='component-box'>"
        html += f"<p><b>Kind of Work:</b> {reading.get('unity_kind_of_work', 'N/A')}</p>"
        html += "<b>Unity Statement:</b>"
        html += f"<div>{reading.get('unity_html', '')}</div>"
        html += "</div>"
        return html

    def _get_terminology_html(self):
        items = self.get_terminology()
        if not items:
            return "<p><i>(No terminology added)</i></p>"
        html = ""
        for item in items:
            term = self.db.get_terminology_details(item['id'])
            html += f"<h3>{term['term']}</h3>"
            html += f"<div>{term.get('meaning', '')}</div>"
        return html

    def _get_propositions_html(self):
        items = self.get_propositions()
        if not items:
            return "<p><i>(No propositions added)</i></p>"
        html = ""
        for item in items:
            prop = self.db.get_proposition_details(item['id'])
            html += f"<h3>{prop['display_name']}</h3>"
            html += f"<div>{prop.get('proposition_html', '')}</div>"
        return html

    # --- TXT Helper Methods ---

    def _get_rubric_txt(self):
        rubric = self.get_rubric()
        if not rubric:
            return "(No rubric components)"
        lines = []
        for item in rubric:
            prefix = "[x] " if item['is_checked'] else "[ ] "
            lines.append(f"{prefix}{item['component_text']}")
        return "\n".join(lines)

    def _get_todo_txt(self):
        items = self.get_todo_list()
        if not items:
            return "(No to-do items)"
        lines = []
        for item in items:
            prefix = "[x] " if item['is_checked'] else "[ ] "
            lines.append(f"{prefix}{item['display_name']}")
        return "\n".join(lines)

    def _get_reading_header_txt(self, reading_id):
        try:
            reading = self.get_reading_details(reading_id)
            text = f"\n---------------------------------\n"
            text += f"READING: {reading['nickname'] or reading['title']}\n"
            text += "---------------------------------\n"
            if reading.get('author'):
                text += f"Author: {reading['author']}\n"
            return text
        except Exception as e:
            return f"(Error processing reading header: {e})"

    def _get_reading_outline_txt(self, reading_id):
        try:
            outline = self.get_reading_outline_data(reading_id)
            if not outline:
                text = "(No outline for this reading)"
                return text

            text = ""
            for item in outline:
                indent = item.get('indent', 0)
                text += f"\n{'  ' * indent}• {item['section_title']}\n"
                notes_html = item.get('notes_html', '')
                if notes_html:
                    text += f"{'  ' * (indent + 1)}Notes: {html_to_text(notes_html)}\n"
            return text
        except Exception as e:
            return f"(Error processing outline: {e})"

    def _get_reading_dqs_txt(self, reading_id):
        items = self.get_reading_dqs(reading_id)
        if not items: return "(No driving questions)\n"
        text = ""
        for item in items:
            text += f"\n{item.get('nickname') or item.get('question_text', '...')}\n"
            text += f"  Why: {html_to_text(item.get('why_question', ''))}\n"
        return text

    def _get_reading_key_terms_txt(self, reading_id):
        items = self.get_reading_key_terms(reading_id)
        if not items: return "(No key terms)\n"
        text = ""
        for item_summary in items:
            item = self.db.get_reading_key_term_details(item_summary['id'])
            text += f"\n{item.get('term', '...')}\n"
            text += f"  Definition: {html_to_text(item.get('definition', ''))}\n"
        return text

    def _get_reading_leading_propositions_txt(self, reading_id):
        items = self.get_reading_leading_propositions(reading_id)
        if not items: return "(No leading propositions)\n"
        text = ""
        for item_summary in items:
            item = self.db.get_reading_proposition_details(item_summary['id'])
            text += f"\n{item.get('nickname') or item.get('proposition_text', '...')}\n"
            text += f"  {html_to_text(item.get('proposition_text', ''))}\n"
        return text

    def _get_reading_arguments_txt(self, reading_id):
        items = self.get_reading_arguments(reading_id)
        if not items: return "(No arguments)\n"
        text = ""
        for item_summary in items:
            item = self.db.get_argument_details(item_summary['id'])
            text += f"\nClaim: {item.get('claim_text', '...')}\n"
            text += f"  Because: {item.get('because_text', '...')}\n"
        return text

    def _get_reading_theories_txt(self, reading_id):
        items = self.get_reading_theories(reading_id)
        if not items: return "(No theories)\n"
        text = ""
        for item_summary in items:
            item = self.db.get_reading_theory_details(item_summary['id'])
            text += f"\n{item.get('theory_name', '...')}\n"
            text += f"  Description: {html_to_text(item.get('description', ''))}\n"
        return text

    def _get_reading_unity_txt(self, reading_id):
        reading = self.get_reading_details(reading_id)
        text = f"Kind of Work: {reading.get('unity_kind_of_work', 'N/A')}\n"
        text += "Unity Statement:\n"
        text += f"{html_to_text(reading.get('unity_html', ''))}\n"
        return text

    def _get_terminology_txt(self):
        items = self.get_terminology()
        if not items:
            return "(No terminology added)"
        text = ""
        for item in items:
            term = self.db.get_terminology_details(item['id'])
            text += f"\n{term['term']}\n"
            text += f"{html_to_text(term.get('meaning', ''))}\n"
        return text

    def _get_propositions_txt(self):
        items = self.get_propositions()
        if not items:
            return "(No propositions added)"
        text = ""
        for item in items:
            prop = self.db.get_proposition_details(item['id'])
            text += f"\n{prop['display_name']}\n"
            text += f"{html_to_text(prop.get('proposition_html', ''))}\n"
        return text